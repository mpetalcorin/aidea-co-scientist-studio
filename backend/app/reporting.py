from io import BytesIO
from typing import Dict, Any

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors


def safe(value: Any, default: str = "") -> str:
    if value is None:
        return default
    return str(value)


def make_docx_report(workspace: Dict[str, Any]) -> BytesIO:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    doc.add_heading("aAidea Co-Scientist Report", 0)

    doc.add_paragraph(
        "Evidence-grounded biomedical hypothesis generation, target intelligence, "
        "pathway mapping, compound retrieval, and experimental planning."
    )

    doc.add_heading("Research Question", level=1)
    doc.add_paragraph(safe(workspace.get("question"), "No research question provided."))

    hypotheses = workspace.get("hypotheses", []) or []
    best = hypotheses[0] if hypotheses else None

    doc.add_heading("Executive Summary", level=1)
    if best:
        doc.add_paragraph(f"Top-ranked hypothesis: {safe(best.get('title'))}")
        doc.add_paragraph(f"Testable version: {safe(best.get('improved'))}")
    else:
        doc.add_paragraph("No top hypothesis available.")

    doc.add_heading("Ranked Hypotheses", level=1)
    if hypotheses:
        for h in hypotheses:
            doc.add_heading(f"{safe(h.get('id'))}: {safe(h.get('title'))}", level=2)
            doc.add_paragraph(f"Rationale: {safe(h.get('rationale'))}")
            doc.add_paragraph(f"Critique: {safe(h.get('critique'))}")
            doc.add_paragraph(f"Improved testable version: {safe(h.get('improved'))}")

            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Metric"
            table.rows[0].cells[1].text = "Score"

            for metric in ["novelty", "plausibility", "feasibility", "translation", "evidence", "risk"]:
                row = table.add_row().cells
                row[0].text = metric.capitalize()
                row[1].text = safe(h.get(metric))
    else:
        doc.add_paragraph("No hypotheses available.")

    doc.add_heading("Target Intelligence", level=1)
    target_intel = workspace.get("targetIntel", []) or []

    if target_intel:
        for item in target_intel:
            summary = item.get("summary", {})
            doc.add_heading(safe(summary.get("target"), "Target"), level=2)
            doc.add_paragraph(safe(summary.get("description")))

            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Field"
            table.rows[0].cells[1].text = "Value"

            rows = [
                ("UniProt available", safe(summary.get("uniprot_available"))),
                ("ChEMBL targets", safe(summary.get("chembl_targets_found"))),
                ("ChEMBL compounds", safe(summary.get("compounds_found"))),
                ("Reactome pathways", safe(summary.get("reactome_pathways_found"))),
                ("Target intelligence score", safe(summary.get("target_intelligence_score"))),
            ]

            for label, value in rows:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = value

            links = []
            if item.get("uniprot", {}).get("url"):
                links.append(f"UniProt: {item['uniprot']['url']}")
            if item.get("chembl", {}).get("targets"):
                first = item["chembl"]["targets"][0]
                if first.get("url"):
                    links.append(f"ChEMBL: {first['url']}")
            if item.get("reactome", {}).get("pathways"):
                first = item["reactome"]["pathways"][0]
                if first.get("url"):
                    links.append(f"Reactome: {first['url']}")

            for link in links:
                doc.add_paragraph(link)
    else:
        doc.add_paragraph("No target intelligence available.")

    doc.add_heading("PubMed Evidence", level=1)
    papers = workspace.get("livePapers", []) or []

    if papers:
        for paper in papers[:12]:
            doc.add_heading(safe(paper.get("title"), "Untitled paper"), level=2)
            doc.add_paragraph(f"Journal: {safe(paper.get('journal'))} {safe(paper.get('year'))}")
            doc.add_paragraph(f"PMID: {safe(paper.get('pmid'), 'N/A')}")
            doc.add_paragraph(safe(paper.get("abstract"), "No abstract retrieved."))
            if paper.get("url"):
                doc.add_paragraph(safe(paper.get("url")))
    else:
        doc.add_paragraph("No PubMed papers retrieved.")

    doc.add_heading("Evidence Audit", level=1)
    evidence_rows = workspace.get("liveEvidence", []) or []

    if evidence_rows:
        for row_item in evidence_rows[:12]:
            doc.add_heading(safe(row_item.get("claim"), "Evidence claim"), level=2)
            doc.add_paragraph(f"Source: {safe(row_item.get('source'))}")
            doc.add_paragraph(f"Strength: {safe(row_item.get('strength'))}")
            doc.add_paragraph(f"Limitation: {safe(row_item.get('limitation'))}")
            if row_item.get("url"):
                doc.add_paragraph(safe(row_item.get("url")))
    else:
        doc.add_paragraph("No evidence audit available.")

    doc.add_heading("Suggested Validation Plan", level=1)
    steps = [
        "Select genetically defined cancer models with appropriate matched controls.",
        "Confirm baseline expression of nominated targets using qPCR, immunoblotting, proteomics, or public omics datasets.",
        "Measure metabolic state using lactate secretion, OCR, ECAR, ATP, NAD+/NADH, ROS, and nucleotide pools.",
        "Test single-agent and combination perturbations, including dose-response and time-course experiments.",
        "Measure DNA damage and replication stress using γH2AX, RAD51 foci, fork progression assays, cell-cycle profiling, and apoptosis markers.",
        "Use computational modelling and feature attribution to connect biomarkers to treatment response.",
        "Prioritise hypotheses with tumour-selective effects, coherent mechanism, and feasible translational path.",
    ]

    for step in steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph(
        "Generated by aAidea Co-Scientist Studio. Automated evidence retrieval requires expert review "
        "before use in grant applications, manuscripts, or clinical decision-making."
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def make_pdf_report(workspace: Dict[str, Any]) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    story = []

    def add_heading(text, level=1):
        style = styles["Title"] if level == 0 else styles["Heading1"] if level == 1 else styles["Heading2"]
        story.append(Paragraph(safe(text), style))
        story.append(Spacer(1, 10))

    def add_para(text):
        story.append(Paragraph(safe(text).replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 8))

    add_heading("aAidea Co-Scientist Report", 0)
    add_para(
        "Evidence-grounded biomedical hypothesis generation, target intelligence, pathway mapping, "
        "compound retrieval, and experimental planning."
    )

    add_heading("Research Question")
    add_para(workspace.get("question", "No research question provided."))

    hypotheses = workspace.get("hypotheses", []) or []
    best = hypotheses[0] if hypotheses else None

    add_heading("Executive Summary")
    if best:
        add_para(f"<b>Top-ranked hypothesis:</b> {safe(best.get('title'))}")
        add_para(f"<b>Testable version:</b> {safe(best.get('improved'))}")
    else:
        add_para("No top hypothesis available.")

    add_heading("Ranked Hypotheses")
    if hypotheses:
        for h in hypotheses:
            add_heading(f"{safe(h.get('id'))}: {safe(h.get('title'))}", 2)
            add_para(f"<b>Rationale:</b> {safe(h.get('rationale'))}")
            add_para(f"<b>Critique:</b> {safe(h.get('critique'))}")
            add_para(f"<b>Improved testable version:</b> {safe(h.get('improved'))}")

            data = [["Metric", "Score"]]
            for metric in ["novelty", "plausibility", "feasibility", "translation", "evidence", "risk"]:
                data.append([metric.capitalize(), safe(h.get(metric))])

            table = Table(data, colWidths=[160, 280])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e0e7ff")),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 14))
    else:
        add_para("No hypotheses available.")

    add_heading("Target Intelligence")
    target_intel = workspace.get("targetIntel", []) or []

    if target_intel:
        for item in target_intel:
            summary = item.get("summary", {})
            add_heading(safe(summary.get("target"), "Target"), 2)
            add_para(safe(summary.get("description")))

            data = [
                ["Field", "Value"],
                ["UniProt available", safe(summary.get("uniprot_available"))],
                ["ChEMBL targets", safe(summary.get("chembl_targets_found"))],
                ["ChEMBL compounds", safe(summary.get("compounds_found"))],
                ["Reactome pathways", safe(summary.get("reactome_pathways_found"))],
                ["Target intelligence score", safe(summary.get("target_intelligence_score"))],
            ]

            table = Table(data, colWidths=[170, 270])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dcfce7")),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 12))
    else:
        add_para("No target intelligence available.")

    add_heading("PubMed Evidence")
    papers = workspace.get("livePapers", []) or []

    if papers:
        for paper in papers[:10]:
            add_heading(safe(paper.get("title"), "Untitled paper"), 2)
            add_para(f"<b>Journal:</b> {safe(paper.get('journal'))} {safe(paper.get('year'))}")
            add_para(f"<b>PMID:</b> {safe(paper.get('pmid'), 'N/A')}")
            add_para(safe(paper.get("abstract"), "No abstract retrieved."))
            if paper.get("url"):
                add_para(safe(paper.get("url")))
    else:
        add_para("No PubMed papers retrieved.")

    add_heading("Suggested Validation Plan")
    validation = [
        "Select genetically defined cancer models with appropriate matched controls.",
        "Confirm baseline expression of nominated targets.",
        "Measure metabolic state using lactate secretion, OCR, ECAR, ATP, NAD+/NADH, ROS, and nucleotide pools.",
        "Test single-agent and combination perturbations.",
        "Measure DNA damage and replication stress.",
        "Use computational modelling and feature attribution.",
        "Prioritise hypotheses with tumour-selective effects and coherent mechanism.",
    ]

    for i, step in enumerate(validation, 1):
        add_para(f"{i}. {step}")

    add_para(
        "Generated by aAidea Co-Scientist Studio. Automated evidence retrieval requires expert review before use."
    )

    doc.build(story)
    buffer.seek(0)
    return buffer
